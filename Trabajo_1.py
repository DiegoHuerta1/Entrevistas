import docx
from docx import Document
import os
import re

# se hace una clase de entrevista
# cada entevista dada sera una instancia de esta clase
# se implementar los atributos y metodos nescesarios para poder
# tratar los datos con facilidad

class Entrevista:

    def __init__(self, nombre_carpeta_doc, name_doc, debug=None ):
        # Se le pasa la locacion del archivo .docx que contiene a la entevista
        self.file_name = name_doc
        # leer este documento, guardar como un atributo
        self.doc = Document('.//'+nombre_carpeta_doc+'//'+name_doc)
        # si debug == 1 entonces se imprimen los inconvenientes
        if debug is not None:
            self.debug = debug
        else:
            self.debug = 0

        # los metadatos de la entrevista: Tipo, sexo, edad, Religion, etc
        # se guardan en un diccionario
        self.metadata = {}

        # los dialogos de la entrevista se guardan en una lista
        # esta lista es de pares (persona, dialogo)
        self.dialogos = []

        # leer la entrevista, obtener los metadatos y los dialogos
        self.leer_entrevista()

        # ya se llenaron self.metadata y self.dialogos
        # poner al entrevistador y entrevistado como atributos
        self.name_entrevistador = self.metadata['entrevistador']
        self.name_entrevistado = self.metadata['seudónimo']
        # obtener el numero de dialogos
        self.num_dialogos = len(self.dialogos)

        # tomar los dialogos y separar en dos listas
        # estas listas son de pares (dialogo, indice en self.dialogos)
        self.dialogos_entrevistador = []
        self.dialogos_entrevistado = []
        self.separar_dialogos()
        self.num_dialogos_entrevistador = len(self.dialogos_entrevistador)
        self.num_dialogos_entrevistado = len(self.dialogos_entrevistado)
        # ver que todo tenga sentido
        assert self.num_dialogos == self.num_dialogos_entrevistador+self.num_dialogos_entrevistado


    def leer_entrevista(self):
        # se usa al inicializar el objeto
        # lee la entrevista, guarda la info importante

        # HIPOTESIS:
            # TODAS LAS ENTERVISTAS ESTAN FORMULADAS COMO:
                # PRIMERO INFORMACION NO RELEVANTE (O REDUNDANTE) (A ESCEPCION DE TALVEZ EL ENTREVISTADOR)
                # DESPUES LA FICHA TECNICA (EMPEZANDO CON UN TEXTO QUE DICE 'Ficha técnica')
                # EN ESTA NO HAY SALTO DE LINEA
                # ADEMAS DE QUE LA INFO ESTA CON EL FORMATO "CAMPO: RESPUESTA"
                # DESPUES UNA LINEA VACIA
                # DESPUES LA ENTREVISTA, UNO A UNO
                # DONDE CADA CAMBIO DE LOCUTOR ES UN NUEVO PARRAFO

        # dividir la entrevista en dos partes:
            # ficha tecnica y entrevista
        self.comenzar_ficha_tecnica = False # sera True cuando se comience a leer
        self.finalizar_ficha_tecnica = False # sera True cuando ya se haya leido


        # iterar en los parrafos de el documento
        for par in self.doc.paragraphs:
            # tomar el texto de este
            txt = par.text

            # ---------------------------------------------------------------------------------#
            # aun no se comienza con la ficha tecnica
            if not(self.comenzar_ficha_tecnica):

                # ver si el parrafo en cuestion es el inicio de la ficha tecnica (titulo)
                if  'Ficha técnica'.lower() in txt.lower():
                    # ya se va a comenzar a leer
                    self.comenzar_ficha_tecnica = True

                # algunas entrevistas tienen al entrevistador antes de la ficha tecnica
                # verdaderamente no se por que, pero pues se adapta el codigo
                elif 'entrevistador' in txt.lower():
                    # se tiene algo como "Entrevistador: persona"
                    partes = txt.split(":", 1) # dividir una vez usando ":"
                    # ver que si se tenga "entrevistador" antes de los dos puntos
                    assert 'entrevistador' in partes[0].strip().lower()
                    entervistador = partes[1].strip() # tomar la segunda parte
                    self.metadata['entrevistador'] = entervistador.strip()

                # haya entrado en algun if o no
                # se sigue a la siguiente iteracion del for
                continue

            # ---------------------------------------------------------------------------------#
            # ya se a comenzado a leer la ficha tecnica, pero no se ha terminado
            if self.comenzar_ficha_tecnica and not(self.finalizar_ficha_tecnica):

                # si se tiene una linea vacia
                if len(txt) == 0:
                    # dos posibilidades

                    # 1) habia espacio enter el titulo y la ficha, seguir buscando
                    # esto sucede cuando la longuitud de los metadatos es a lo mas 1
                    # pues se podria tener al entrevistador de antes
                    if len(list(self.metadata.keys())) <= 1:
                        continue

                    # 2) ya se llego al final de la ficha tecnica
                    else:
                        self.finalizar_ficha_tecnica = True
                        continue

                # nuestro texto no es vacio
                # se asume que se tiene algo del tipo "campo:  respuesta"
                partes = txt.split(":", 1) # dividir una vez usando ":"

                # si no se puede dividir, entonces no esta en el formato que se busca
                # se termina entonces de leer la ficha tecnica
                if len(partes) == 1:
                    self.finalizar_ficha_tecnica = True
                    continue

                # si se pudo dividir
                campo = partes[0].strip()
                respuesta = partes[1].strip()
                self.metadata[campo.lower().strip()] = respuesta.strip()
                continue
            # ---------------------------------------------------------------------------------#
            # ya se termino de leer la ficha tecnica

            # se esta leyendo la entrevista tal cual
            if self.finalizar_ficha_tecnica:

                # si es linea vacia, se salta
                if len(txt) == 0:
                    continue

                # aveces hay cosas como [comentario]
                # estos no se quieren considerar
                if txt[0].strip() == '[':
                    if self.debug == 1:
                        print('Se hace caso omiso al comentario:')
                        print(txt)
                    continue

                # aca se tienen cosas del estilo "nombre- texto", "nombre: texto"
                # primero se separa el nombre del texto
                patron = r"[:\-;\s+]|(?<=\w)[.]"  # separacion
                coincidencias = re.split(patron, txt, maxsplit=1) # solo separar una vez

                # si es diferente de dos, no se pudo separar
                # entonces solo se ignora este parrafo, pasar al siguiente
                if len(coincidencias) != 2:
                    if self.debug == 1:
                        print('Texto con formato inseperado:')
                        print(txt)
                    continue

                # si se pudo separar
                nombre = coincidencias[0].strip()
                texto = coincidencias[1].strip()

                # ver si el texto es diferente a vacio
                # en ese caso se omite
                if len(texto) == 0:
                    if self.debug == 1:
                        print('Dialogo sin contenido')
                        print(txt)
                    continue

                # ver quien es el que esta hablando, y agregar el dialogo a las lista
                if nombre.lower() in self.metadata['entrevistador'].lower():
                    # esta hablando el entrevistador
                    self.dialogos.append(('entrevistador', texto))
                elif nombre.lower() in self.metadata['seudónimo'].lower():
                    # esta hablando el entrevistado
                    self.dialogos.append(('entrevistado', texto))
                else:
                    # no puedo saber quien esta hablando
                    if self.debug == 1:
                        print("No se identifico si este dialogo es del entervistador o entrevistato (se omite)")
                        print(txt)
                        print('Nombre identificado:', nombre)
                    continue
            # --------------------------------------------------------------------------------

    def separar_dialogos(self):
        # ya se tiene lleno el self.dialogos
        # separar en self.dialogos_entrevistador y self.dialogos_entrevistado
        # que son listas de pares (texto, indice)

        # iterar los dialogos
        for idx, par in enumerate(self.dialogos):
            persona = par[0]
            texto = par[1]
            # es el entrevistador
            if persona == 'entrevistador':
                self.dialogos_entrevistador.append((texto, idx))
            # es el entrevistado
            elif persona == 'entrevistado':
                self.dialogos_entrevistado.append((texto, idx))
            else:
                raise Exception("Dialogo no es de entrevistador ni entrevistado")


    # devuelve los dialogos de la entrevista, es una lista de pares (persona, dialogo)
    def get_dialogos(self, incluir_persona = False):
        # si incluir_persona == True, entonces se devuelve completa. (persona, dialogo)
        if incluir_persona:
            return self.dialogos
        # si incluir_persona == False, entonces solo se devuelven los dialogos (dialogo)
        elif incluir_persona == False:
            return [dialogo[1] for dialogo in self.dialogos]


    # devuelve los dialogos del entrevistador, es una lista de pares (texto, indice)
    def get_dialogos_entrevistador(self, incluir_indice = False):
        # si incluir_indice == True, entonces se devuelve completa. (texto, indice)
        if incluir_indice:
            return self.dialogos_entrevistador
        # si incluir_indice == False, entonces solo se devuelven los textos (texto)
        elif incluir_indice == False:
            return [par[0] for par in self.dialogos_entrevistador]

    # devuelve las preguntas que hizo el entrevistador
    # esto es, los dialogos del entrevistador que contengan: '¿' o '?'
    def get_preguntas(self):
        dialogos_entrevistador = self.get_dialogos_entrevistador()
        preguntas = [dialogo for dialogo in dialogos_entrevistador if '¿' in dialogo or '?' in dialogo]
        return preguntas

    # devuelve los dialogos del entrevistado, es una lista de pares (texto, indice)
    def get_dialogos_entrevistado(self, incluir_indice = False):
        # si incluir_indice == True, entonces se devuelve completa. (texto, indice)
        if incluir_indice:
            return self.dialogos_entrevistado
        # si incluir_indice == False, entonces solo se devuelven los textos (texto)
        elif incluir_indice == False:
            return [par[0] for par in self.dialogos_entrevistado]

    # devuele el nombre del file de la entrevista
    def get_name_doc(self):
        return self.file_name[:-5] # no dar el .docx

    # devuelve los metadatos, dict de pares (campo, valor)
    def get_metadatos(self):
        return self.metadata

    # devuelve los campos de los metadatos
    def get_tipos_metadatos(self):
        return list(self.metadata.keys())

    # devuelve el numero de metadatos
    def get_num_metadatos(self):
        return len(list(self.metadata.keys()))

    # devolver todo el texto
    def get_full_text(self):
        full_text = []
        # iterar los parrafos
        for par in self.doc.paragraphs:
            # agregar el texto, como elemento de una lista
            full_text.append(par.text)
        # pasar de lista a string
        full_string = '.\n'.join(full_text)
        return full_string

    # dar un mini resumen de la entrevista
    def print_resumen(self):
        print("-"*100)
        print(self.file_name)
        print('Numero de datos en la ficha tecnica:', self.get_num_metadatos())
        print('Entrevistador:', self.name_entrevistador)
        print('Entrevistado:', self.name_entrevistado)
        print('Numero de dialogos del entrevistador:', self.num_dialogos_entrevistador)
        print('Numero de dialogos del entrevistado:', self.num_dialogos_entrevistado)

        print("-"*100)

    # imprime los metadatos con buen formato
    def print_metadatos(self):
        # Iterar sobre pares clave-valor
        for campo, valor in self.metadata.items():
            print(f'{campo}: {valor}')

    # imprimir toda la info importante
    def print_todo(self):
        print("-"*100)
        print(self.file_name)
        print(" ")
        self.print_metadatos()
        print(" ")
        print('Numero de dialogos:', self.num_dialogos)
        for dialogo in self.dialogos:
            print(" ")
            print(f'{dialogo[0]}: {dialogo[1]}')

        print("-"*100)
